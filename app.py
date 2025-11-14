#one edit for lasting 6
#label app - AIS
import streamlit as st
import pandas as pd
from docx import Document
import barcode
from barcode.writer import ImageWriter
from zipfile import ZipFile
import io
import os
import hashlib
#one edit for lasting 3

# Set page config
st.set_page_config(page_title="Label & Barcode Generator - AIS", layout="centered", page_icon="📦")


st.markdown(
    "<h1 style='color:#023E8A;text-align:center;'>📦 Label & Barcode Generator - AIS</h1>",
    unsafe_allow_html=True
)

# Step 1: Table input
st.subheader("Step 1: Enter Packaging Details")
num_rows = st.number_input("How many rows do you want to input?", min_value=1, max_value=1000, value=1)

table_data = {
    "No.": ["" for _ in range(num_rows)],
    "Part Number": ["" for _ in range(num_rows)],
    "Description": ["" for _ in range(num_rows)],
    "Quantity per pack": ["" for _ in range(num_rows)],
    "Number of packs": ["" for _ in range(num_rows)]
}

df = pd.DataFrame(table_data)
edited_df = st.data_editor(df, num_rows="dynamic", use_container_width=True)

# Step 2: Additional inputs
st.subheader("Step 2: Enter Additional Information")
customer = st.text_input("Customer Name")
po_number = st.text_input("PO Number")
mfg = st.text_input("Mfg. (e.g. 03-2025)")

if st.button("Generate Packaging Labels"):
    # Validate inputs
    if edited_df.isnull().values.any() or not customer or not po_number or not mfg:
        st.error("Please fill in all the fields before generating labels.")
    else:
        # Create in-memory zip file
        zip_buffer = io.BytesIO()
        
        with ZipFile(zip_buffer, 'w') as zipf:
            # Create and add Word document directly to zip
            doc = Document()
            
            for index, row in edited_df.iterrows():
                generated_part_numbers = set()
                no = str(row['No.']).strip()
                part_number = str(row['Part Number']).strip()
                description = str(row['Description']).strip()
                qty_pack = str(row['Quantity per pack']).strip()
                num_packs = str(row['Number of packs']).strip()

                # Add to Word
                doc.add_paragraph(f"{no}. {customer} {part_number} Pack Label")
                doc.add_paragraph(f"Part Number: {part_number}")
                doc.add_paragraph(f"Description: {description}")
                doc.add_paragraph(f"Quantity: {qty_pack} PCS")
                doc.add_paragraph(f"PO Number: {po_number}")
                doc.add_paragraph("Bar Code:")
                doc.add_paragraph(f"Mfg. {mfg}")
                doc.add_paragraph(f"Quantity: {num_packs}")
                doc.add_page_break()

                # Generate barcode directly in memory
                if part_number:
                    barcode_filename = f"{no}_{part_number}.png"

                    # Skip if the filename contains ".2."
                    if ".2" in barcode_filename:
                        continue

                    barcode_buffer = io.BytesIO()
                    code128 = barcode.get("code128", part_number, writer=ImageWriter())
                    code128.write(barcode_buffer)
                    zipf.writestr(f"barcodes/{barcode_filename}", barcode_buffer.getvalue())



            # Save Word doc directly to zip
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            zipf.writestr("Packaging_Labels.docx", doc_buffer.getvalue())

        # Prepare download
        zip_buffer.seek(0)
        st.success("✅ Labels and Barcodes Generated!")
        st.download_button(
            "📥 Download All Files (ZIP)",
            data=zip_buffer,
            file_name="Packaging_Labels.zip",
            mime="application/zip"
        )
